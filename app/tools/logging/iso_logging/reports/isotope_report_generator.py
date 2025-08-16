"""
同位素解释报告生成工具 - 收集工具执行结果并生成Word格式报告
"""
import os
import re
import json
import time
import logging
from datetime import datetime
from typing import Dict, List, Any, Optional, Union, Tuple
import uuid
from pathlib import Path

import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT, WD_BREAK
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn  # 导入命名空间，用于设置东亚字体
from io import BytesIO
import matplotlib.pyplot as plt
from PIL import Image

from langchain_core.tools import BaseTool
from langchain_core.messages import BaseMessage, HumanMessage, SystemMessage
from langchain_openai import ChatOpenAI
from langchain_core.language_models.llms import BaseLLM

from app.tools.registry import register_tool
from app.utils.qwen_chat import SFChatOpenAI
from app.core.state import IsotopeSystemState, TaskStatus

# 配置日志
logger = logging.getLogger(__name__)

def extract_image_path(content):
    """从输出内容中提取图片路径
    
    Args:
        content: 输出内容
        
    Returns:
        图片路径，如果未找到则返回None
    """
    patterns = [
        # 处理原始格式的图片路径
        r'(文件|图片|image|file)(?:路径|path)[:：]\s*([^\s,\n\(\)]+\.(png|jpg|jpeg|svg|pdf))',
        r'(?:保存|saved)(?:到|to|at)[:：]?\s*([^\s,\n\(\)]+\.(png|jpg|jpeg|svg|pdf))',
        r'路径[:：]?\s*([^\s,\n\(\)]+\.(png|jpg|jpeg|svg|pdf))',
        r'(?:生成|保存)(?:了|的)(?:图表|图像|图片|图形|可视化结果|文件)[:：]?\s*([^\s,\n\(\)]+\.(png|jpg|jpeg|svg|pdf))',
        r'(?:图表|图像|图片|图形|可视化结果|文件)已(?:生成|保存)(?:在|到)[:：]?\s*([^\s,\n\(\)]+\.(png|jpg|jpeg|svg|pdf))',
        
        # Markdown图片链接格式
        r'!\[.*?\]\(([^\s\(\)]+\.(png|jpg|jpeg|svg|pdf))\)',
        
        # 处理带引号的路径
        r'[\'"]([^\'"\s]+\.(png|jpg|jpeg|svg|pdf))[\'"]',
        
        # 处理Windows路径（带双反斜杠）
        r'([A-Za-z]:\\\\[^\\]+(?:\\\\[^\\]+)*\.(png|jpg|jpeg|svg|pdf))',
        r'([A-Za-z]:\\[^\\]+(?:\\[^\\]+)*\.(png|jpg|jpeg|svg|pdf))',
        
        # 处理JSON对象中的file_path字段（单引号）
        r'file_path\':\s*\'([^\']+\.(png|jpg|jpeg|svg|pdf))\'',
        
        # 处理JSON对象中的file_path字段（双引号）
        r'file_path":\s*"([^"]+\.(png|jpg|jpeg|svg|pdf))"',
        
        # 处理文件ID行中包含的路径
        r'文件ID:.*?file_path[\'"]?:\s*[\'"]([^\'"]+\.(png|jpg|jpeg|svg|pdf))[\'"]',
        
        # 处理不同格式的文件路径引用
        r'[\'"]?file_path[\'"]?\s*[:=]\s*[\'"]([^\'"]+\.(png|jpg|jpeg|svg|pdf))[\'"]',
        
        # 处理包含在JSON对象中的file_name字段
        r'[\'"]?file_name[\'"]?\s*[:=]\s*[\'"]([^\'"]+\.(png|jpg|jpeg|svg|pdf))[\'"]',
    ]
    
    # 尝试所有模式
    for pattern in patterns:
        match = re.search(pattern, content, re.IGNORECASE)
        if match:
            # 根据模式的不同，组索引可能是1或2
            path = match.group(1) if len(match.groups()) == 1 or pattern.startswith(r'!') or pattern.startswith(r'[\'"]') else match.group(2)
            # 处理Windows路径中的双反斜杠
            path = path.replace('\\\\', '\\')
            
            logger.info(f"使用模式 '{pattern}' 提取到图片路径: {path}")
            if os.path.exists(path):
                return path
            else:
                logger.warning(f"提取到的路径不存在: {path}")
    
    # 尝试查找可能的JSON对象描述的图片路径
    try:
        # 查找类似 {'file_path': '...'} 的模式
        json_pattern = r'\{[^{}]*?[\'"]file_path[\'"][^{}]*?\}'
        match = re.search(json_pattern, content)
        if match:
            json_str = match.group(0)
            # 尝试将单引号形式转换为JSON可解析的双引号形式
            try:
                fixed_json = json_str.replace("'", '"')
                fixed_json = re.sub(r'(\w+):', r'"\1":', fixed_json)  # 确保所有键都有引号
                json_obj = json.loads(fixed_json)
                
                if "file_path" in json_obj:
                    path = json_obj["file_path"].replace('\\\\', '\\')
                    if os.path.exists(path):
                        logger.info(f"从JSON对象中提取到图片路径: {path}")
                        return path
            except Exception as e:
                logger.warning(f"解析可能的JSON对象失败: {str(e)}")
    except Exception as e:
        logger.warning(f"尝试解析JSON对象时出错: {str(e)}")
    
    # 尝试提取包含文件ID的完整JSON对象
    try:
        # 匹配包含文件ID和JSON对象的完整模式
        file_id_pattern = r'文件ID:\s*(\{[^}]+\})'
        match = re.search(file_id_pattern, content)
        if match:
            json_str = match.group(1)
            # 修复JSON格式
            try:
                fixed_json = json_str.replace("'", '"')
                fixed_json = re.sub(r'(\w+):', r'"\1":', fixed_json)  # 确保所有键都有引号
                json_obj = json.loads(fixed_json)
                
                # 尝试从不同字段获取路径
                if "file_path" in json_obj:
                    path = json_obj["file_path"].replace('\\\\', '\\')
                    if os.path.exists(path):
                        logger.info(f"从文件ID的JSON对象中提取到file_path: {path}")
                        return path
                elif "file_name" in json_obj:
                    # 尝试构建完整路径
                    file_name = json_obj["file_name"]
                    possible_paths = [
                        os.path.join("data", "generated", file_name),
                        os.path.join("data", file_name),
                        file_name
                    ]
                    for p in possible_paths:
                        if os.path.exists(p):
                            logger.info(f"根据file_name构建并验证路径: {p}")
                            return p
            except Exception as e:
                logger.warning(f"解析文件ID的JSON对象失败: {str(e)}")
    except Exception as e:
        logger.warning(f"尝试提取文件ID对象时出错: {str(e)}")
    
    # 尝试从更广泛的JSON对象文本中提取
    try:
        # 先尝试匹配大段JSON格式的文本
        json_blocks = re.findall(r'\{[^{]*?file_(?:path|name)[^}]*?\}', content)
        for block in json_blocks:
            try:
                # 修复JSON格式
                fixed_json = block.replace("'", '"')
                fixed_json = re.sub(r'(\w+):', r'"\1":', fixed_json)  # 确保所有键都有引号
                # 尝试解析
                json_obj = json.loads(fixed_json)
                
                # 尝试提取路径
                if "file_path" in json_obj:
                    path = json_obj["file_path"].replace('\\\\', '\\')
                    if os.path.exists(path):
                        logger.info(f"从广泛JSON块中提取到file_path: {path}")
                        return path
                elif "file_name" in json_obj:
                    # 尝试在常见目录中构建完整路径
                    file_name = json_obj["file_name"]
                    data_dir = os.path.join("data", "generated")
                    if os.path.exists(data_dir):
                        for f in os.listdir(data_dir):
                            if f == file_name or file_name in f:
                                path = os.path.join(data_dir, f)
                                logger.info(f"从generated目录找到匹配文件: {path}")
                                return path
            except Exception as e:
                logger.debug(f"解析单个JSON块失败: {str(e)}")
    except Exception as e:
        logger.warning(f"尝试匹配大段JSON文本时出错: {str(e)}")
    
    return None

def format_document_title(doc, title, level=1):
    """格式化文档标题
    
    Args:
        doc: 文档对象
        title: 标题文本
        level: 标题级别(1-4)
    """
    if level == 1:
        # 主标题
        heading = doc.add_heading(title, level=0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in heading.runs:
            run.font.size = Pt(18)
            run.font.bold = True
            run.font.name = 'SimHei'  # 西文字体设为黑体
            # 设置东亚字体为黑体
            try:
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
            except Exception as e:
                logger.warning(f"设置东亚字体失败: {str(e)}")
            run.font.color.rgb = RGBColor(0, 0, 128)  # 深蓝色
    elif level == 2:
        # 章节标题
        heading = doc.add_heading(title, level=1)
        for run in heading.runs:
            run.font.size = Pt(16)
            run.font.bold = True
            run.font.name = 'SimHei'  # 西文字体设为黑体
            # 设置东亚字体为黑体
            try:
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
            except Exception as e:
                logger.warning(f"设置东亚字体失败: {str(e)}")
            run.font.color.rgb = RGBColor(0, 64, 128)  # 蓝色
    elif level == 3:
        # 子章节标题
        heading = doc.add_heading(title, level=2)
        for run in heading.runs:
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.name = 'SimHei'  # 西文字体设为黑体
            # 设置东亚字体为黑体
            try:
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
            except Exception as e:
                logger.warning(f"设置东亚字体失败: {str(e)}")
            run.font.color.rgb = RGBColor(30, 100, 130)  # 青蓝色
    else:
        # 普通标题
        heading = doc.add_heading(title, level=3)
        for run in heading.runs:
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.name = 'SimHei'  # 西文字体设为黑体
            # 设置东亚字体为黑体
            try:
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
            except Exception as e:
                logger.warning(f"设置东亚字体失败: {str(e)}")

def add_image_to_document(doc, image_path, caption=None, width=6):
    """向文档中添加图片
    
    Args:
        doc: 文档对象
        image_path: 图片路径
        caption: 图片说明
        width: 图片宽度(英寸)
    """
    try:
        # 检查文件是否存在
        if not os.path.exists(image_path):
            logger.warning(f"图片文件不存在: {image_path}")
            add_paragraph_with_style(doc, f"[图片加载失败: {image_path}]")
            return False
        
        # 添加图片
        doc.add_picture(image_path, width=Inches(width))
        
        # 添加图片说明
        if caption:
            cap_para = doc.add_paragraph(caption)
            cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap_para.style = 'Caption'
            # 确保图片说明文字使用宋体
            for run in cap_para.runs:
                run.font.name = 'SimSun'  # 西文字体设为宋体
                # 设置东亚字体为宋体
                try:
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
                except Exception as e:
                    logger.warning(f"设置东亚字体失败: {str(e)}")
            
        return True
    except Exception as e:
        logger.error(f"添加图片失败: {str(e)}")
        add_paragraph_with_style(doc, f"[图片添加失败: {str(e)}]")
        return False

def create_report_document():
    """创建新的报告文档，设置样式
    
    Returns:
        docx文档对象
    """
    doc = Document()
    
    # 设置文档属性
    doc.core_properties.title = "天然气碳同位素数据解释报告"
    doc.core_properties.author = "同位素智能体系统"
    doc.core_properties.created = datetime.now()
    
    # 设置默认字体为宋体（同时设置西文和东亚字体）
    style = doc.styles['Normal']
    font = style.font
    font.name = 'SimSun'  # 西文字体设为宋体
    # 设置东亚字体为宋体
    try:
        style._element.get_or_add_rPr().rFonts.set(qn('w:eastAsia'), 'SimSun')
    except Exception as e:
        logger.warning(f"设置东亚字体失败: {str(e)}")
    font.size = Pt(12)
    
    # 设置所有标题样式的字体为黑体
    for i in range(1, 10):
        if f'Heading {i}' in doc.styles:
            heading_style = doc.styles[f'Heading {i}']
            heading_style.font.name = 'SimHei'  # 西文字体设为黑体
            # 设置东亚字体为黑体
            try:
                heading_style._element.get_or_add_rPr().rFonts.set(qn('w:eastAsia'), 'SimHei')
            except Exception as e:
                logger.warning(f"设置标题东亚字体失败: {str(e)}")
    
    # 创建图片说明样式
    if 'Caption' not in doc.styles:
        caption_style = doc.styles.add_style('Caption', WD_STYLE_TYPE.PARAGRAPH)
        caption_style.font.italic = True
        caption_style.font.size = Pt(10)
        caption_style.font.name = 'SimSun'  # 西文字体设为宋体
        # 设置东亚字体为宋体
        try:
            caption_style._element.get_or_add_rPr().rFonts.set(qn('w:eastAsia'), 'SimSun')
        except Exception as e:
            logger.warning(f"设置图片说明东亚字体失败: {str(e)}")
        caption_style.font.color.rgb = RGBColor(100, 100, 100)  # 灰色
    
    return doc

def generate_chapter_content(
    llm: BaseLLM,
    chapter_title: str,
    tool_results: List[Dict[str, Any]],
    include_images: bool = True
) -> Tuple[str, List[str]]:
    """生成章节内容
    
    Args:
        llm: 语言模型
        chapter_title: 章节标题
        tool_results: 章节相关的工具执行结果列表
        include_images: 是否包含图片
        
    Returns:
        章节内容和图片路径列表
    """
    logger.info(f"开始生成章节 '{chapter_title}' 的内容，工具结果数量: {len(tool_results)}")
    
    results_text = []
    image_paths = []
    
    # 处理每个工具执行结果
    for i, result in enumerate(tool_results):
        try:
            # 提取工具输出内容
            content = result.get("output", "")
            if not content and "content" in result:
                content = result["content"]
            
            if not isinstance(content, str):
                # 处理非字符串类型的内容（如字典、列表等）
                content = str(content)
            
            logger.info(f"处理第 {i+1}/{len(tool_results)} 个工具结果，内容长度: {len(content)}")
            results_text.append(content)
            
            # 提取图片路径
            if include_images:
                # 首先尝试从result的特定字段中提取
                if "file_path" in result:
                    path = result["file_path"]
                    if os.path.exists(path):
                        logger.info(f"从结果字段中直接找到图片路径: {path}")
                        image_paths.append(path)
                        continue
                        
                # 然后尝试从内容中提取
                path = extract_image_path(content)
                if path:
                    logger.info(f"从结果内容中提取到图片路径: {path}")
                    image_paths.append(path)
                    continue
                    
                # 尝试解析JSON格式的内容
                try:
                    # 查找可能包含文件路径的JSON字符串
                    json_pattern = r'\{[^{}]*?file_path[^{}]*?\}'
                    matches = re.finditer(json_pattern, content)
                    
                    for match in matches:
                        json_str = match.group(0)
                        try:
                            # 尝试修复常见的JSON格式问题
                            fixed_json = json_str.replace("'", '"')
                            fixed_json = re.sub(r'(\w+):', r'"\1":', fixed_json)  # 确保所有键都有引号
                            json_obj = json.loads(fixed_json)
                            
                            if "file_path" in json_obj:
                                path = json_obj["file_path"].replace('\\\\', '\\')
                                if os.path.exists(path):
                                    logger.info(f"从JSON对象中提取到图片路径: {path}")
                                    image_paths.append(path)
                                    break
                        except Exception as e:
                            logger.warning(f"解析JSON字符串失败: {str(e)}, 字符串: {json_str}")
                except Exception as e:
                    logger.warning(f"尝试提取JSON格式的图片路径时出错: {str(e)}")
        
        except Exception as e:
            logger.error(f"处理工具结果时出错: {str(e)}")
            results_text.append(f"处理结果出错: {str(e)}")
    
    logger.info(f"为章节 '{chapter_title}' 提取到 {len(image_paths)} 个图片路径")
    
    # 构建系统提示
    system_prompt = f"""你是一位专业地质学家和技术报告撰写专家。请根据以下工具执行结果，撰写一个关于"{chapter_title}"的专业报告章节。
要求：
1. 内容要专业、科学，围绕章节主题展开
2. 使用严谨的学术语言，避免口语化表达
3. 内容应包含工具执行结果中的关键数据和发现
4. 文字应当连贯、有条理，不要简单罗列工具执行结果
5. 不要编造不存在于工具执行结果中的数据或结论
6. 保留所有重要的数值数据和分析结果
7. 不要提及AI、模型、提示词等内容
8. 仅关注与{chapter_title}相关的内容

请输出格式良好的Markdown文本。"""

    # 调用语言模型生成内容
    messages = [
        SystemMessage(content=system_prompt),
        HumanMessage(content="\n\n".join(results_text))
    ]
    
    try:
        response = llm.invoke(messages)
        content = response.content
    except Exception as e:
        logger.error(f"生成章节内容时出错: {str(e)}")
        content = f"## {chapter_title}\n\n生成内容时出错，请重试。错误信息: {str(e)}"
    
    return content, image_paths

def collect_tool_results(state):
    """从系统状态中收集工具执行结果
    
    Args:
        state: 系统状态或状态字典
        
    Returns:
        工具执行结果列表，按类型分组
    """
    # 初始化结果集合
    categorized_results = {
        "数据加载": [],
        "数据分析": [],
        "数据可视化": [],
        "同位素分类": [],
        "深度趋势": [],
        "综合解释": [],
        "其他": []
    }
    
    # 处理不同类型的状态对象
    if not state:
        logger.warning("状态为空，无法收集工具执行结果")
        return categorized_results
        
    logger.info(f"状态类型: {type(state)}")
    
    # 获取当前任务
    current_task = None
    try:
        # 尝试作为字典访问
        if isinstance(state, dict):
            current_task = state.get("current_task", {})
        # 尝试作为对象访问
        elif hasattr(state, "get"):
            current_task = state.get("current_task", {})
    except Exception as e:
        logger.error(f"获取当前任务失败: {str(e)}")
        return categorized_results
    
    if not current_task:
        logger.warning("当前任务为空，无法收集工具执行结果")
        return categorized_results
        
    logger.info(f"当前任务类型: {type(current_task)}")
    
    # 获取工具执行结果
    tool_executions = []
    try:
        # 尝试作为字典访问
        if isinstance(current_task, dict):
            tool_executions = current_task.get("tool_executions", [])
        # 尝试作为对象访问
        elif hasattr(current_task, "get"):
            tool_executions = current_task.get("tool_executions", [])
    except Exception as e:
        logger.error(f"获取工具执行结果失败: {str(e)}")
        return categorized_results
    
    logger.info(f"收集到 {len(tool_executions)} 个工具执行结果")
    
    # 分类工具执行结果
    for execution in tool_executions:
        try:
            if not isinstance(execution, dict) and not hasattr(execution, "get"):
                logger.warning(f"跳过非字典工具执行记录: {type(execution)}")
                continue
                
            # 获取工具名称
            tool_name = ""
            if isinstance(execution, dict):
                tool_name = execution.get("tool_name", "")
            elif hasattr(execution, "get"):
                tool_name = execution.get("tool_name", "")
                
            if not tool_name:
                logger.warning("工具执行记录中无工具名称，跳过")
                continue
                
            # 根据工具名称进行分类
            if "load" in tool_name.lower() or "read" in tool_name.lower() or "读取" in tool_name:
                categorized_results["数据加载"].append(execution)
            elif "plot" in tool_name.lower() or "visualize" in tool_name.lower() or "图" in tool_name:
                categorized_results["数据可视化"].append(execution)
            elif "classify" in tool_name.lower() or "分类" in tool_name or "类别" in tool_name:
                categorized_results["同位素分类"].append(execution)
            elif "depth" in tool_name.lower() or "trend" in tool_name.lower() or "深度" in tool_name:
                categorized_results["深度趋势"].append(execution)
            elif "analyze" in tool_name.lower() or "analysis" in tool_name.lower() or "分析" in tool_name:
                categorized_results["数据分析"].append(execution)
            elif "interpret" in tool_name.lower() or "interpretation" in tool_name.lower() or "解释" in tool_name:
                categorized_results["综合解释"].append(execution)
            else:
                categorized_results["其他"].append(execution)
        except Exception as e:
            logger.error(f"处理工具执行记录时出错: {str(e)}")
            continue
    
    return categorized_results

def add_paragraph_with_style(doc, text, style_name='Normal'):
    """添加带有指定样式的段落
    
    Args:
        doc: 文档对象
        text: 段落文本
        style_name: 样式名称
        
    Returns:
        添加的段落对象
    """
    paragraph = doc.add_paragraph(text)
    paragraph.style = style_name
    
    # 确保段落中的所有文本都使用宋体（同时设置西文和东亚字体）
    for run in paragraph.runs:
        run.font.name = 'SimSun'  # 西文字体设为宋体
        # 设置东亚字体为宋体
        try:
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
        except Exception as e:
            logger.warning(f"设置东亚字体失败: {str(e)}")
    
    return paragraph

def add_empty_line(doc):
    """添加空行
    
    Args:
        doc: 文档对象
        
    Returns:
        添加的段落对象
    """
    para = doc.add_paragraph()
    para.style = 'Normal'
    return para

@register_tool(category="iso_logging", use_structured_tool=True)
def generate_isotope_report(state_dict: Dict[str, Any] = None) -> str:
    """生成天然气碳同位素数据解释综合报告，收集当前任务中所有工具执行结果，并生成docx格式报告
    
    Args:
        state_dict: None
        
    Returns:
        生成的报告文件路径
    """
    # 获取流写入器，用于向前端推送执行过程
    from langgraph.config import get_stream_writer
    writer = get_stream_writer()
    
    if writer:
        writer({"custom_step": "开始生成天然气碳同位素数据解释综合报告..."})
    
    # 创建输出目录
    output_dir = os.path.join("data", "generated")
    os.makedirs(output_dir, exist_ok=True)
    
    if writer:
        writer({"custom_step": "创建输出目录成功"})
    
    # 生成报告文件路径
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    report_filename = f"isotope_report_{timestamp}.docx"
    output_path = os.path.join(output_dir, report_filename)
    
    if writer:
        writer({"custom_step": f"报告文件将保存为: {report_filename}"})
    
    # 初始化LLM
    try:
        llm = SFChatOpenAI(temperature=0.3)
        logger.info("成功初始化SFChatOpenAI")
        if writer:
            writer({"custom_step": "成功初始化智能语言模型"})
    except Exception as e:
        logger.warning(f"初始化SFChatOpenAI失败: {str(e)}，使用备用LLM")
        if writer:
            writer({"custom_step": f"初始化首选模型失败，切换到备用模型: {str(e)}"})
        llm = ChatOpenAI(temperature=0.3)
    
    # 创建文档
    doc = create_report_document()
    if writer:
        writer({"custom_step": "创建报告文档模板成功"})
    
    # 添加报告标题
    format_document_title(doc, "天然气碳同位素数据解释综合报告", level=1)
    
    # 添加报告生成时间
    add_paragraph_with_style(doc, f"报告生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    add_paragraph_with_style(doc, "由同位素智能体自动生成")
    add_empty_line(doc)
    
    # 处理状态
    system_state = state_dict if isinstance(state_dict, dict) else {}
    if not system_state:
        logger.warning("未提供系统状态，生成的报告可能不完整")
        if writer:
            writer({"custom_step": "警告: 未提供完整系统状态，报告可能不完整"})
        system_state = {}
    
    logger.info(f"接收到的状态字典键: {list(system_state.keys())}")
    
    # 收集并分类工具执行结果
    if writer:
        writer({"custom_step": "正在收集并分类工具执行结果..."})
    
    categorized_results = collect_tool_results(system_state)
    
    # 打印分类后的结果数量
    result_summary = []
    for category, results in categorized_results.items():
        logger.info(f"分类 '{category}' 包含 {len(results)} 个工具结果")
        if len(results) > 0:
            result_summary.append(f"{category}: {len(results)}个")
    
    if writer and result_summary:
        writer({"custom_step": f"收集到的工具结果: {', '.join(result_summary)}"})
    
    # 定义报告章节结构
    report_chapters = [
        {"title": "1. 概述", "key": "综合解释", "include_images": True},
        {"title": "2. 数据来源与方法", "key": "数据加载", "include_images": False},
        {"title": "3. 同位素分布特征", "key": "数据可视化", "include_images": True},
        {"title": "4. 天然气类型分类", "key": "同位素分类", "include_images": True},
        {"title": "5. 深度变化趋势分析", "key": "深度趋势", "include_images": True},
        {"title": "6. 地质意义解释", "key": "综合解释", "include_images": False},
        {"title": "7. 结论与建议", "key": "综合解释", "include_images": False}
    ]
    
    if writer:
        writer({"custom_step": f"报告将包含{len(report_chapters)}个章节，开始生成内容..."})
    
    # 生成各章节内容
    for chapter in report_chapters:
        if writer:
            writer({"custom_step": f"正在生成章节: {chapter['title']}..."})
        
        # 添加章节标题
        format_document_title(doc, chapter["title"], level=2)
        
        # 获取该章节相关的工具结果
        chapter_results = categorized_results.get(chapter["key"], [])
        if chapter["key"] == "综合解释" and not chapter_results:
            # 如果没有特定综合解释结果，使用所有结果
            chapter_results = []
            for key, results in categorized_results.items():
                chapter_results.extend(results)
        
        if not chapter_results:
            # 如果没有相关结果，添加占位内容
            add_paragraph_with_style(doc, "本章节暂无相关数据和分析结果。")
            if writer:
                writer({"custom_step": f"章节 {chapter['title']} 无相关数据和分析结果"})
            continue
        
        # 使用LLM生成章节内容
        if writer:
            writer({"custom_step": f"分析{len(chapter_results)}个工具结果，为章节 '{chapter['title']}' 生成内容..."})
        
        content, image_paths = generate_chapter_content(
            llm, 
            chapter["title"], 
            chapter_results, 
            include_images=chapter["include_images"]
        )
        
        if writer:
            writer({"custom_step": f"章节 '{chapter['title']}' 内容生成完成，提取到{len(image_paths)}张相关图片"})
        
        # 添加生成的内容
        paragraphs = content.split('\n\n')
        for para in paragraphs:
            if para.strip():
                add_paragraph_with_style(doc, para.strip())
        
        # 添加图片
        if chapter["include_images"] and image_paths:
            if writer:
                writer({"custom_step": f"正在添加{len(image_paths)}张图片到章节 '{chapter['title']}'..."})
            
            add_empty_line(doc)
            for idx, img_path in enumerate(image_paths):
                add_image_to_document(
                    doc, 
                    img_path, 
                    caption=f"图 {chapter['title'].split('.')[0]}-{idx+1}: {os.path.basename(img_path)}"
                )
                add_empty_line(doc)
    
    # 保存文档
    try:
        if writer:
            writer({"custom_step": "报告内容生成完成，正在保存文档..."})
        
        # 确保输出路径存在
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        
        doc.save(output_path)
        logger.info(f"成功生成报告: {output_path}")
        
        if writer:
            writer({"custom_step": f"报告成功保存: {output_path}"})
        
        # 检查文件是否确实创建
        if os.path.exists(output_path):
            file_size = os.path.getsize(output_path)
            logger.info(f"报告文件大小: {file_size} 字节")
            
            if writer:
                writer({"custom_step": f"报告文件大小: {file_size} 字节"})
            
            # 准备返回文件消息
            file_message = {
                "file_path": output_path,
                "file_name": report_filename,
                "file_type": "docx"
            }
            
            if writer:
                writer({"file_message": file_message})
            
            # 返回成功消息和文件信息
            return json.dumps({
                "status": "success",
                "message": f"报告生成成功，文件路径: {output_path}",
                "file_message": file_message
            })
        else:
            error_msg = f"文件保存成功但找不到文件: {output_path}"
            logger.error(error_msg)
            if writer:
                writer({"custom_step": f"错误: {error_msg}"})
            return json.dumps({
                "status": "error",
                "message": error_msg
            })
    except Exception as e:
        error_msg = f"保存报告失败: {str(e)}"
        logger.error(error_msg)
        if writer:
            writer({"custom_step": f"错误: {error_msg}"})
        return json.dumps({
            "status": "error",
            "message": error_msg
        }) 